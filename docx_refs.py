#!/usr/bin/env python3

import docx
import os
import re
import sys

doc_path = sys.argv[1]

try:
    ref_section_name = sys.argv[2]
except IndexError:
    # Default if user did not provide one
    ref_section_name = 'References'

try:
    tablefig_section_name = sys.argv[3]
except IndexError:
    tablefig_section_name = 'Tables and Figures'

# User provided blank string
if not ref_section_name:
    ref_section_name = 'References'

if not tablefig_section_name:
    tablefig_section_name = 'Tables and Figures'

try:
    doc = docx.Document(doc_path)
except docx.opc.exceptions.PackageNotFoundError:
    print('Not a valid .docx file.')
    exit()

# List of error messages for main Golang server
err = []

# Find location of references section heading and table/fig section
# heading
for i, par in enumerate(doc.paragraphs):
    if par.text.strip() == ref_section_name:
        ref_loc = i
        break

for i, par in enumerate(doc.paragraphs):
    if par.text.strip() == tablefig_section_name:
        tablefig_loc = i
        break

# Add to error messages if not found
try:
    ref_loc
except NameError:
    err.append(('"{}" section not found.'.format(ref_section_name)))

try:
    tablefig_loc
except NameError:
    err.append(('"{}" section not found.'.format(tablefig_section_name)))

# Exit now if sections not found. Printed messages will be read to
# Golang server.
if err:
    for msg in err:
        print(msg)
    exit()

# Extract authors and years
par = [x.text for x in doc.paragraphs]
refs = par[ref_loc+1:tablefig_loc]
authors = []
years = []
for ref in refs:
    # Extract 'first author' field
    try:
        authors.append(re.split('\. |, ', ref)[0])
    except IndexError:
        err.append('One or more references have an invalid author field.')
    # Extract 'year' field
    try:
        years.append(re.findall('[1-3][0-9]{3}', ref)[0])
    except IndexError:
        err.append('One or more references have an invalid year field.')

# Exit if references are invalid
if err:
    err = list(set(err))
    for msg in err:
        print(msg)
    exit()

# Alphabetize
ref_zip = list(zip(doc.paragraphs[ref_loc+1:tablefig_loc],
                   refs, authors, years))

# Sort by author (x[2]) then year (x[3])
ref_zip = sorted(ref_zip, key=lambda x: (x[2], x[3]))

# Write out text files of original and new references:
orig_refs = [x.text for x in doc.paragraphs[ref_loc+1:tablefig_loc]]
sorted_refs = [x[0].text for x in ref_zip]

with open('{}_orig'.format(doc_path), 'w+') as f:
    for ref in orig_refs:
        f.write('{}\n\n'.format(ref))

with open('{}_sorted'.format(doc_path), 'w+') as f:
    for ref in sorted_refs:
        f.write('{}\n\n'.format(ref))

# Git diff the two files
os.system('git diff --color-words --no-index --no-prefix -U1000 {}_orig {}_sorted | tail -n +6 | ansi2html > {}'.format(doc_path, doc_path, doc_path))
